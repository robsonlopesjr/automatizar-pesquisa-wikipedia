import wikipedia as wp
from docx import Document

wp.set_lang('pt')

assunto = input("Sobre o que você deseja pesquisar?\n")

while True:
    try:
        wiki = wp.page(assunto)
        break
    except Exception:
        print("Assunto inválido!")
        assunto = input("Sobre o que você deseja pesquisar?\n")

texto = wiki.content

word = Document()
titulo = word.add_paragraph(assunto)
titulo.aligment = 1
texto_word = word.add_paragraph(texto)
texto_word.aligment = 2
word.save(assunto + ".docx")
